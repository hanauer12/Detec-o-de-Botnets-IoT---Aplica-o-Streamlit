#!/usr/bin/env python3
"""
Script para converter texto_projeto_sbc.md para formato DOCX
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def parse_markdown_to_docx(md_file, docx_file):
    """Converte arquivo Markdown para DOCX formatado"""
    
    # Lê o arquivo Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Cria documento Word
    doc = Document()
    
    # Configuração de estilo padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Título principal (#)
        if line.startswith('# ') and not line.startswith('##'):
            title_text = line[2:].strip()
            title = doc.add_heading(title_text, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Subtítulos (##)
        if line.startswith('## '):
            subtitle_text = line[3:].strip()
            # Remove --- se houver
            if subtitle_text == '---':
                i += 1
                continue
            doc.add_heading(subtitle_text, level=2)
            i += 1
            continue
        
        # Subtítulos nível 3 (###)
        if line.startswith('### '):
            subtitle_text = line[4:].strip()
            doc.add_heading(subtitle_text, level=3)
            i += 1
            continue
        
        # Subtítulos nível 4 (####)
        if line.startswith('#### '):
            subtitle_text = line[5:].strip()
            doc.add_heading(subtitle_text, level=4)
            i += 1
            continue
        
        # Linha horizontal (---)
        if line == '---':
            i += 1
            continue
        
        # Texto em negrito (**texto**)
        if '**' in line:
            # Processa negrito
            parts = re.split(r'(\*\*.*?\*\*)', line)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.strip():
                    p.add_run(part)
            i += 1
            continue
        
        # Tabela (formato Markdown simples)
        if '|' in line and i < len(lines) - 1:
            # Detecta início de tabela
            table_lines = []
            j = i
            while j < len(lines) and '|' in lines[j].strip():
                table_lines.append(lines[j].strip())
                j += 1
            
            if len(table_lines) >= 2:
                # Cria tabela
                header_line = table_lines[0]
                separator_line = table_lines[1]  # Linha de separação
                
                # Parse do header
                headers = [h.strip() for h in header_line.split('|') if h.strip()]
                
                # Remove linha de separação se for apenas ---
                data_start = 2 if separator_line.replace('|', '').replace('-', '').replace(':', '').strip() == '' else 1
                
                # Conta linhas de dados
                data_lines = table_lines[data_start:]
                num_rows = len([l for l in data_lines if l.strip() and not all(c in '-|: ' for c in l)])
                
                if headers and num_rows > 0:
                    table = doc.add_table(rows=num_rows + 1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Header
                    header_cells = table.rows[0].cells
                    for idx, header in enumerate(headers):
                        header_cells[idx].text = header
                        for paragraph in header_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # Dados
                    row_idx = 1
                    for data_line in data_lines:
                        if data_line.strip() and not all(c in '-|: ' for c in data_line):
                            cells = [c.strip() for c in data_line.split('|') if c.strip()]
                            if len(cells) == len(headers):
                                for col_idx, cell_text in enumerate(cells):
                                    table.rows[row_idx].cells[col_idx].text = cell_text
                                row_idx += 1
                
                i = j
                continue
        
        # Código inline (`texto`)
        if '`' in line:
            parts = re.split(r'(`[^`]+`)', line)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                elif part.strip():
                    p.add_run(part)
            i += 1
            continue
        
        # Lista não ordenada (- ou *)
        if line.startswith('- ') or line.startswith('* '):
            items = []
            j = i
            while j < len(lines) and (lines[j].strip().startswith('- ') or lines[j].strip().startswith('* ') or lines[j].strip() == ''):
                if lines[j].strip():
                    items.append(lines[j].strip()[2:])
                j += 1
            
            for item in items:
                p = doc.add_paragraph(item, style='List Bullet')
            
            i = j
            continue
        
        # Lista numerada
        if re.match(r'^\d+\.\s', line):
            items = []
            j = i
            while j < len(lines) and (re.match(r'^\d+\.\s', lines[j].strip()) or lines[j].strip() == ''):
                if lines[j].strip():
                    items.append(re.sub(r'^\d+\.\s', '', lines[j].strip()))
                j += 1
            
            for item in items:
                p = doc.add_paragraph(item, style='List Number')
            
            i = j
            continue
        
        # Parágrafo normal
        if line:
            # Remove formatação markdown básica
            clean_line = line
            # Remove links [texto](url) -> texto
            clean_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_line)
            doc.add_paragraph(clean_line)
        
        i += 1
    
    # Salva documento
    doc.save(docx_file)
    print(f"✅ Documento Word criado: {docx_file}")

if __name__ == '__main__':
    try:
        parse_markdown_to_docx('texto_projeto_sbc.md', 'texto_projeto_sbc.docx')
    except ImportError:
        print("❌ Erro: python-docx não está instalado.")
        print("💡 Instale com: pip install python-docx")
        print("   Ou use: pip3 install python-docx")
    except Exception as e:
        print(f"❌ Erro ao converter: {e}")
        import traceback
        traceback.print_exc()



